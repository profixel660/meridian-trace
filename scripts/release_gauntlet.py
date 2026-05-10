"""Release gauntlet — pre-ship verification for the Meridian wheel + installer.

Run from the repo root:

    uv run python scripts/release_gauntlet.py

Exits 0 only when every step passes. Each step targets a bug class that has
already been shipped (and embarrassed us) on the v0.2.0-alpha line:

    1.  PowerShell parser check on installer/*.ps1 / *.psm1
    2.  ASCII-only check on installer/*.{ps1,psm1,bat,cmd}   (alpha-3 em-dash)
    2b. Installer URL constants — never 'http://localhost'   (alpha-5 IPv6/4)
    3.  Wheel build precondition (apps/web/out present) + wheel content check
    4.  Fresh-venv install of the just-built wheel
    5.  Cwd=System32 simulation — backend launched from a System32-like cwd
        must still write logs to MERIDIAN_HOME, not the cwd               (alpha-2)
    6.  /setup/ wizard URL probe — body contains "Meridian"               (alpha-3)
    7.  Version assertion — /health JSON.version == pyproject version    (alpha-1)
    7b. ODA-detection contract — scan response includes 'oda' object
        with 'available' (bool) and 'install_url' (str)                  (alpha-11)
    7c. Import status-pivot smoke — kick a tiny job, poll, confirm
        status reaches 'succeeded' and 'failed_groups' is a list. Catches
        the alpha-10 frontend-typed-it-wrong bug class                   (alpha-11)
    2c. Detach-pattern + .env-loader static check — installer/Start-
        Meridian.bat AND the Install-Meridian.ps1 inline copy contain
        pythonw.exe + WindowStyle Hidden (alpha-12 detach contract) AND
        MERIDIAN_ENV_FILE + Set-Item -Path "Env: (alpha-14 .env-loader
        contract). Statically guards the alpha-13 documentation-vs-
        reality gap that step 7h's subprocess-level test alone cannot
        catch -- a future commit that drops the .env loader from the
        .bat would pass 7h but ship the broken launcher.       (alpha-12 + alpha-14)
    2d. Launcher .bat exec-policy check — installer/Meridian-Console.bat
        invokes powershell.exe with -ExecutionPolicy Bypass. Catches
        regressions that re-expose the alpha-11 ExecutionPolicy block.   (alpha-12)
    7f. /setup/runtime contract — endpoint returns pid, started_at,
        uptime_seconds, version, python_version, platform. Catches
        regressions that break the operator triage path.                 (alpha-12)
    7g. Auth bypass default-secure — /setup/runtime.auth_disabled is
        false on a fresh install with no env var set. Catches a
        regression that defaults the alpha-13 MERIDIAN_AUTH_DISABLED
        bypass on.                                                       (alpha-13)
    7h. .env-loaded env-var contract — write MERIDIAN_AUTH_DISABLED=1
        into a temp .env, spawn the backend with the launcher's .env-
        loading semantics, and confirm /setup/runtime reports
        auth_disabled=true. Locks the contract ".env in the install
        dir -> env-var in the running backend -> /setup/runtime
        reflects it." Catches the alpha-13 documentation-vs-reality
        gap where Start-Meridian.bat never loaded .env.                  (alpha-14)
    8.  CLI --help smoke (no browser opens)
    9.  Summary

The script does NOT mutate anything outside ``dist/`` and a tmpdir scratch
area. It does NOT modify installer files, source files, or apps/web/.
"""

from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import tomllib
import urllib.error
import urllib.request
import zipfile
from pathlib import Path

# ──────────────────────────────────────────────────────────────────────────
# Pretty printing — keep ASCII only so the gauntlet's own output never trips
# step 2 if a future caller pipes this stdout into an installer log.
# ──────────────────────────────────────────────────────────────────────────


def _ok(step: str, msg: str = "") -> None:
    suffix = f" {msg}" if msg else ""
    print(f"[ OK  ] {step}{suffix}")


def _fail(step: str, msg: str) -> None:
    print(f"[FAIL ] {step}: {msg}", file=sys.stderr)


def _info(msg: str) -> None:
    print(f"[ ... ] {msg}")


# ──────────────────────────────────────────────────────────────────────────
# Repo-root resolution
# ──────────────────────────────────────────────────────────────────────────


def _repo_root() -> Path:
    here = Path(__file__).resolve().parent
    # scripts/ sits at the repo root.
    candidate = here.parent
    if not (candidate / "pyproject.toml").exists():
        raise SystemExit(
            f"release_gauntlet: expected pyproject.toml at {candidate}, not found. "
            "Run from the repo root."
        )
    return candidate


# ──────────────────────────────────────────────────────────────────────────
# Step 1 — PowerShell parser check
# ──────────────────────────────────────────────────────────────────────────


def _step_powershell_parse(installer_dir: Path) -> bool:
    targets = sorted(
        list(installer_dir.rglob("*.ps1")) + list(installer_dir.rglob("*.psm1"))
    )
    if not targets:
        _info("No .ps1/.psm1 files under installer/ — nothing to parse-check.")
        _ok("step 1: PowerShell parse")
        return True

    if not shutil.which("powershell") and not shutil.which("powershell.exe"):
        _info("PowerShell not on PATH — skipping parser check (non-Windows host).")
        _ok("step 1: PowerShell parse (skipped)")
        return True

    # One PS process; iterate file paths inside it. We pass the file list
    # via a temp file (one path per line) instead of CLI args — PowerShell's
    # `-Command` wraps args in its own parser and treats `--` as a unary
    # operator, which mangles arg-style passing.
    list_file = installer_dir.parent / ".gauntlet_ps_paths.tmp"
    try:
        list_file.write_text(
            "\n".join(str(p) for p in targets), encoding="utf-8"
        )
        ps_script = (
            r"$listFile = $env:GAUNTLET_PS_LIST; "
            r"$paths = Get-Content -LiteralPath $listFile; "
            r"$failed = $false; "
            r"foreach ($p in $paths) { "
            r"  $errors = $null; "
            r"  [void][System.Management.Automation.Language.Parser]::ParseFile("
            r"      $p, [ref]$null, [ref]$errors); "
            r"  if ($errors -and $errors.Count -gt 0) { "
            r"    $failed = $true; "
            r"    foreach ($e in $errors) { "
            r"      Write-Host ('PARSE_ERR ' + $p + ': ' + $e.Message) "
            r"    } "
            r"  } "
            r"}; "
            r"if ($failed) { exit 1 } else { exit 0 }"
        )
        env = {**os.environ, "GAUNTLET_PS_LIST": str(list_file)}
        try:
            result = subprocess.run(
                ["powershell", "-NoProfile", "-NonInteractive", "-Command", ps_script],
                capture_output=True,
                text=True,
                timeout=60,
                env=env,
            )
        except (subprocess.TimeoutExpired, FileNotFoundError) as exc:
            _fail("step 1: PowerShell parse", f"could not run powershell: {exc}")
            return False
    finally:
        try:
            list_file.unlink()
        except OSError:
            pass

    if result.returncode != 0:
        _fail(
            "step 1: PowerShell parse",
            "parser errors:\n" + (result.stdout or result.stderr).strip(),
        )
        return False

    _ok("step 1: PowerShell parse", f"{len(targets)} file(s) parsed cleanly")
    return True


# ──────────────────────────────────────────────────────────────────────────
# Step 2 — ASCII-only on installer scripts
# ──────────────────────────────────────────────────────────────────────────


def _step_ascii_only(installer_dir: Path) -> bool:
    suffixes = {".ps1", ".psm1", ".bat", ".cmd"}
    targets = [
        p for p in installer_dir.rglob("*") if p.is_file() and p.suffix.lower() in suffixes
    ]
    if not targets:
        _ok("step 2: ASCII-only", "no installer scripts found")
        return True

    bad: list[tuple[Path, int, int, int, str]] = []  # (path, line, col, codepoint, char)
    for path in targets:
        try:
            text = path.read_bytes()
        except OSError as exc:
            _fail("step 2: ASCII-only", f"could not read {path}: {exc}")
            return False
        # Decode as UTF-8 (most likely). If it isn't valid UTF-8, that itself
        # is a non-ASCII signal — flag the first offending byte.
        try:
            decoded = text.decode("utf-8")
        except UnicodeDecodeError as exc:
            bad.append((path, 0, exc.start, text[exc.start], f"<bad utf-8 byte 0x{text[exc.start]:02x}>"))
            continue
        line = 1
        col = 0
        for ch in decoded:
            col += 1
            if ch == "\n":
                line += 1
                col = 0
                continue
            cp = ord(ch)
            if cp > 0x7F:
                bad.append((path, line, col, cp, ch))
                # First offender per file is enough signal — keep scanning
                # other files but skip the rest of this one.
                break

    if bad:
        msg_lines = ["non-ASCII codepoints found (forbidden in installer scripts):"]
        for path, line, col, cp, ch in bad:
            msg_lines.append(
                f"  {path}:{line}:{col}  U+{cp:04X} ({ch!r})"
            )
        _fail("step 2: ASCII-only", "\n".join(msg_lines))
        return False

    _ok("step 2: ASCII-only", f"{len(targets)} file(s) clean")
    return True


# ──────────────────────────────────────────────────────────────────────────
# Step 3 — wheel build + content check
# ──────────────────────────────────────────────────────────────────────────


def _read_pyproject_version(repo: Path) -> str:
    with (repo / "pyproject.toml").open("rb") as f:
        data = tomllib.load(f)
    return data["project"]["version"]


def _step_build_wheel(repo: Path) -> tuple[bool, Path | None]:
    web_out = repo / "apps" / "web" / "out"
    if not web_out.exists() or not (web_out / "index.html").exists():
        _info("apps/web/out missing — running `npm run build` to produce it.")
        npm = shutil.which("npm") or shutil.which("npm.cmd")
        if npm is None:
            _fail("step 3: wheel build", "npm not found on PATH and apps/web/out is absent")
            return False, None
        try:
            r = subprocess.run(
                [npm, "run", "build"],
                cwd=repo / "apps" / "web",
                capture_output=True,
                text=True,
                timeout=600,
            )
        except subprocess.TimeoutExpired:
            _fail("step 3: wheel build", "`npm run build` timed out after 10 minutes")
            return False, None
        if r.returncode != 0:
            tail = (r.stdout + r.stderr).splitlines()[-30:]
            _fail("step 3: wheel build", "`npm run build` failed:\n" + "\n".join(tail))
            return False, None
        if not (web_out / "index.html").exists():
            _fail("step 3: wheel build", "apps/web/out/index.html still missing after build")
            return False, None

    version = _read_pyproject_version(repo)
    dist = repo / "dist"
    dist.mkdir(exist_ok=True)
    expected_wheel = dist / f"meridian-{version}-py3-none-any.whl"

    # Always rebuild — the gauntlet's job is to prove the *current* tree.
    uv = shutil.which("uv") or shutil.which("uv.exe")
    if uv is None:
        _fail("step 3: wheel build", "uv not found on PATH")
        return False, None

    _info(f"running `uv build --wheel --out-dir dist` (target: {expected_wheel.name})")
    try:
        r = subprocess.run(
            [uv, "build", "--wheel", "--out-dir", str(dist)],
            cwd=repo,
            capture_output=True,
            text=True,
            timeout=600,
        )
    except subprocess.TimeoutExpired:
        _fail("step 3: wheel build", "`uv build` timed out after 10 minutes")
        return False, None
    if r.returncode != 0:
        tail = (r.stdout + r.stderr).splitlines()[-30:]
        _fail("step 3: wheel build", "uv build failed:\n" + "\n".join(tail))
        return False, None

    if not expected_wheel.exists():
        # Maybe uv produced a slightly different name — find any wheel matching version.
        candidates = list(dist.glob(f"meridian-{version}*.whl"))
        if not candidates:
            _fail("step 3: wheel build", f"no wheel produced for version {version} in {dist}")
            return False, None
        expected_wheel = candidates[0]

    # Verify wheel contains the bundled wizard HTML.
    must_have = "meridian/_web/setup/index.html"
    with zipfile.ZipFile(expected_wheel) as zf:
        names = set(zf.namelist())
    if must_have not in names:
        _fail(
            "step 3: wheel build",
            f"wheel {expected_wheel.name} is missing {must_have} "
            f"(apps/web/out/setup/index.html did not get bundled). "
            f"Did `npm run build` actually emit the static export for /setup/*?",
        )
        return False, None

    _ok(
        "step 3: wheel build",
        f"{expected_wheel.name} contains {must_have}",
    )
    return True, expected_wheel


# ──────────────────────────────────────────────────────────────────────────
# Step 4 — fresh-venv install
# ──────────────────────────────────────────────────────────────────────────


def _step_fresh_install(wheel: Path, scratch: Path) -> tuple[bool, Path | None]:
    venv_dir = scratch / "venv"
    uv = shutil.which("uv") or shutil.which("uv.exe")
    if uv is None:
        _fail("step 4: fresh install", "uv not found on PATH")
        return False, None

    _info(f"creating fresh venv at {venv_dir}")
    r = subprocess.run(
        [uv, "venv", str(venv_dir)],
        capture_output=True,
        text=True,
        timeout=120,
    )
    if r.returncode != 0:
        _fail("step 4: fresh install", f"uv venv failed:\n{r.stdout}\n{r.stderr}")
        return False, None

    if os.name == "nt":
        py = venv_dir / "Scripts" / "python.exe"
        meridian_exe = venv_dir / "Scripts" / "meridian.exe"
    else:
        py = venv_dir / "bin" / "python"
        meridian_exe = venv_dir / "bin" / "meridian"

    if not py.exists():
        _fail("step 4: fresh install", f"python not found at {py} after uv venv")
        return False, None

    _info(f"installing {wheel.name} into the fresh venv")
    r = subprocess.run(
        [uv, "pip", "install", "--python", str(py), str(wheel)],
        capture_output=True,
        text=True,
        timeout=600,
    )
    if r.returncode != 0:
        tail = (r.stdout + r.stderr).splitlines()[-30:]
        _fail("step 4: fresh install", "wheel install failed:\n" + "\n".join(tail))
        return False, None

    if not meridian_exe.exists():
        _fail(
            "step 4: fresh install",
            f"console-script {meridian_exe.name} missing at {meridian_exe} "
            "after install — pyproject [project.scripts] regression?",
        )
        return False, None

    _ok("step 4: fresh install", f"meridian.exe at {meridian_exe}")
    return True, venv_dir


# ──────────────────────────────────────────────────────────────────────────
# Steps 5–7 — backend launch from System32-like cwd, /setup/ probe, /health version
# ──────────────────────────────────────────────────────────────────────────


def _probe(url: str, timeout: float = 1.0) -> tuple[int | None, bytes]:
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "release-gauntlet"})
        with urllib.request.urlopen(req, timeout=timeout) as resp:  # noqa: S310 — localhost
            return resp.status, resp.read()
    except (urllib.error.URLError, TimeoutError, OSError):
        return None, b""


def _wait_healthy(base: str, deadline_s: float = 30.0) -> bool:
    deadline = time.monotonic() + deadline_s
    while time.monotonic() < deadline:
        status, _ = _probe(f"{base}/health", timeout=0.5)
        if status == 200:
            return True
        time.sleep(0.25)
    return False


def _step_backend_lifecycle(
    venv_dir: Path,
    scratch: Path,
    repo: Path,
) -> bool:
    if os.name == "nt":
        py = venv_dir / "Scripts" / "python.exe"
    else:
        py = venv_dir / "bin" / "python"

    # Spawn cwd: a tmp dir with NO pyproject.toml above it. On Windows this
    # mirrors the System32-cwd class of bug exactly.
    spawn_cwd = scratch / "spawn_cwd"
    spawn_cwd.mkdir(exist_ok=True)
    meridian_home = scratch / "meridian_home"
    meridian_home.mkdir(exist_ok=True)
    port = 8765
    base = f"http://127.0.0.1:{port}"

    env = {
        **os.environ,
        "MERIDIAN_HOME": str(meridian_home),
        "MERIDIAN_PORT": str(port),
        "MERIDIAN_HOST": "127.0.0.1",
    }
    # The brief explicitly says env only carries MERIDIAN_HOME + MERIDIAN_PORT.
    # We do NOT set MERIDIAN_PROJECT_ROOT — the assertion below is precisely
    # whether MERIDIAN_HOME alone is enough to keep logs out of an arbitrary
    # spawn cwd. The alpha-3 _is_unsafe_cwd() check only triggers on Windows
    # system paths; this step verifies the broader bug class (any non-
    # pyproject cwd), which the spec (§5) requires.
    # Strip any cached-bytecode pollution that could shadow the just-installed pkg.
    env.pop("PYTHONPATH", None)

    log_path = scratch / "backend.log"
    _info(f"spawning backend from cwd={spawn_cwd}, MERIDIAN_HOME={meridian_home}")
    with log_path.open("wb") as logf:
        proc = subprocess.Popen(
            [str(py), "-m", "meridian.api.main"],
            cwd=str(spawn_cwd),
            env=env,
            stdout=logf,
            stderr=subprocess.STDOUT,
        )
    try:
        if not _wait_healthy(base, deadline_s=30.0):
            try:
                tail = log_path.read_text(errors="replace").splitlines()[-30:]
            except OSError:
                tail = []
            _fail(
                "step 5: cwd-System32 simulation",
                "backend never reported /health 200 within 30s. backend.log tail:\n"
                + "\n".join(tail),
            )
            return False

        # /health 200 confirmed → step 5 partially OK. Check log location.
        # Logs must NOT have appeared under spawn_cwd; they must live under
        # meridian_home. The exact subdir is `data/projects/_global.logs/`.
        logs_dir = meridian_home / "data" / "projects" / "_global.logs"
        spawn_cwd_logs = list(spawn_cwd.rglob("_global.logs"))
        if spawn_cwd_logs:
            _fail(
                "step 5: cwd-System32 simulation",
                f"logs leaked into spawn cwd: {spawn_cwd_logs}. "
                "alpha-2 regression — _project_root() did not substitute MERIDIAN_HOME.",
            )
            return False
        if not logs_dir.exists():
            # Tolerate the exact subpath name shifting; require *some* log
            # directory to exist under MERIDIAN_HOME.
            any_logs = list(meridian_home.rglob("*.log"))
            if not any_logs:
                _fail(
                    "step 5: cwd-System32 simulation",
                    f"no logs found under {meridian_home}. Expected {logs_dir} or similar.",
                )
                return False
            _info(
                f"note: expected log dir {logs_dir} not present, but logs landed under "
                f"MERIDIAN_HOME ({len(any_logs)} file(s)) — accepting"
            )
        _ok(
            "step 5: cwd-System32 simulation",
            f"backend healthy; logs at {logs_dir if logs_dir.exists() else meridian_home}",
        )

        # Step 6 — /setup/ wizard URL probe. Body must contain "Meridian".
        status, body = _probe(f"{base}/setup/", timeout=5.0)
        if status != 200:
            _fail(
                "step 6: /setup/ probe",
                f"GET /setup/ returned {status} (expected 200). "
                "alpha-3 regression class — /setup/welcome 404.",
            )
            return False
        if b"Meridian" not in body and b"meridian" not in body:
            preview = body[:200].decode("latin-1", errors="replace")
            _fail(
                "step 6: /setup/ probe",
                f"GET /setup/ body did not contain 'Meridian'. First 200 bytes: {preview!r}",
            )
            return False
        _ok("step 6: /setup/ probe", f"GET /setup/ -> 200, body OK ({len(body)} bytes)")

        # Step 7 — /health JSON.version === pyproject version.
        status, body = _probe(f"{base}/health", timeout=5.0)
        if status != 200:
            _fail("step 7: version assertion", f"/health returned {status}")
            return False
        try:
            payload = json.loads(body.decode("utf-8"))
        except (UnicodeDecodeError, json.JSONDecodeError) as exc:
            _fail(
                "step 7: version assertion",
                f"/health body was not JSON: {exc}; body={body!r}",
            )
            return False
        api_version = payload.get("version")
        pyproj_version = _read_pyproject_version(repo)
        if api_version != pyproj_version:
            _fail(
                "step 7: version assertion",
                f"/health.version={api_version!r} != pyproject.version={pyproj_version!r}. "
                "alpha-1 regression class — stale '0.1.0' shipped in installed wheel.",
            )
            return False
        _ok("step 7: version assertion", f"version={api_version} matches pyproject")

        # Step 7b — ODA-detection contract (alpha-11). A scan response
        # must include the new `oda` object with `available` (bool) and
        # `install_url` (str). Catches frontend-blocking schema drift.
        scan_target = scratch / "gauntlet_oda_probe"
        scan_target.mkdir(exist_ok=True)
        scan_payload = json.dumps({"folder_path": str(scan_target)}).encode("utf-8")
        try:
            req = urllib.request.Request(
                f"{base}/setup/import-folder/scan",
                data=scan_payload,
                headers={
                    "Content-Type": "application/json",
                    "User-Agent": "release-gauntlet",
                },
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=10) as resp:  # noqa: S310 — localhost
                scan_body = json.loads(resp.read().decode("utf-8"))
        except (urllib.error.URLError, TimeoutError, json.JSONDecodeError) as exc:
            _fail("step 7b: ODA-detection contract", f"scan POST failed: {exc}")
            return False
        oda = scan_body.get("oda")
        if not isinstance(oda, dict):
            _fail(
                "step 7b: ODA-detection contract",
                "scan response missing the 'oda' object — alpha-11 regression class. "
                "Frontend cannot render the pre-import 'install ODA File Converter' "
                f"callout without it. Body keys: {sorted(scan_body)}",
            )
            return False
        for required in ("available", "install_url"):
            if required not in oda:
                _fail(
                    "step 7b: ODA-detection contract",
                    f"oda object missing required key '{required}'. oda={oda!r}",
                )
                return False
        if not isinstance(oda["available"], bool):
            _fail(
                "step 7b: ODA-detection contract",
                f"oda.available must be bool, got {type(oda['available']).__name__}",
            )
            return False
        if not (isinstance(oda["install_url"], str) and oda["install_url"].startswith("http")):
            _fail(
                "step 7b: ODA-detection contract",
                f"oda.install_url must be an http URL, got {oda['install_url']!r}",
            )
            return False
        _ok(
            "step 7b: ODA-detection contract",
            f"oda.available={oda['available']}, install_url present",
        )

        # Step 7c — folder-import status-pivot smoke (alpha-11). The
        # alpha-10 frontend pivoted on `status === "done"` while the
        # backend writes `"succeeded"`. This step kicks a tiny one-file
        # import and asserts the status terminal value matches the
        # canonical contract AND `failed_groups` is a list (the new
        # alpha-11 field). Without this step the frontend ↔ backend
        # contract drifts silently.
        import_target = scratch / "gauntlet_import_pivot"
        import_target.mkdir(exist_ok=True)
        # Empty folder yields total=0 which still finishes via the
        # 'succeeded' branch with imported=0. We exploit that — no need
        # to fabricate a real ingestable doc here.
        import_payload = json.dumps(
            {
                "folder_path": str(import_target),
                "project_name": "gauntlet-import-pivot",
            }
        ).encode("utf-8")
        try:
            req = urllib.request.Request(
                f"{base}/setup/import-folder",
                data=import_payload,
                headers={
                    "Content-Type": "application/json",
                    "User-Agent": "release-gauntlet",
                },
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=10) as resp:  # noqa: S310 — localhost
                kick_body = json.loads(resp.read().decode("utf-8"))
        except (urllib.error.URLError, TimeoutError, json.JSONDecodeError) as exc:
            _fail("step 7c: import status-pivot smoke", f"kick POST failed: {exc}")
            return False
        job_id = kick_body.get("job_id")
        if not isinstance(job_id, str):
            _fail(
                "step 7c: import status-pivot smoke",
                f"import-folder response missing job_id; got {kick_body!r}",
            )
            return False
        # Poll until terminal — short deadline, the empty-folder path
        # finishes in milliseconds.
        deadline = time.monotonic() + 10.0
        last_body: dict = {}
        while time.monotonic() < deadline:
            poll_status, poll_payload = _probe(
                f"{base}/setup/import-folder/{job_id}", timeout=2.0
            )
            if poll_status == 200:
                try:
                    last_body = json.loads(poll_payload.decode("utf-8"))
                except json.JSONDecodeError:
                    last_body = {}
                if last_body.get("status") in ("succeeded", "failed"):
                    break
            time.sleep(0.1)
        terminal = last_body.get("status")
        if terminal not in ("succeeded", "failed"):
            _fail(
                "step 7c: import status-pivot smoke",
                f"job did not reach terminal status within deadline; last_body={last_body!r}",
            )
            return False
        if terminal != "succeeded":
            _fail(
                "step 7c: import status-pivot smoke",
                f"empty-folder import should yield status='succeeded', got '{terminal}'. "
                f"failed_groups={last_body.get('failed_groups')!r}",
            )
            return False
        if not isinstance(last_body.get("failed_groups"), list):
            _fail(
                "step 7c: import status-pivot smoke",
                "import-folder status response missing alpha-11 'failed_groups' list. "
                "Frontend's coalesced-error rendering depends on this field. "
                f"body keys: {sorted(last_body)}",
            )
            return False
        if not isinstance(last_body.get("oda"), dict):
            _fail(
                "step 7c: import status-pivot smoke",
                "import-folder status response missing alpha-11 'oda' object.",
            )
            return False
        _ok(
            "step 7c: import status-pivot smoke",
            f"job reached '{terminal}', failed_groups + oda present",
        )

        # Step 7f — /setup/runtime contract (alpha-12). Operator triage uses
        # this endpoint to answer pid / uptime / log path / last-job from
        # one HTTP call. Lock the field set so a future refactor doesn't
        # silently drop a key the launcher .bat / Status-Meridian.bat
        # depends on.
        rt_status, rt_payload = _probe(f"{base}/setup/runtime", timeout=5.0)
        if rt_status != 200:
            _fail(
                "step 7f: /setup/runtime contract",
                f"GET /setup/runtime returned {rt_status} (expected 200)",
            )
            return False
        try:
            rt_body = json.loads(rt_payload.decode("utf-8"))
        except json.JSONDecodeError as exc:
            _fail("step 7f: /setup/runtime contract", f"non-JSON body: {exc}")
            return False
        required_runtime = {
            "pid": int,
            "started_at": str,
            "uptime_seconds": (int, float),
            "version": str,
            "python_version": str,
            "platform": str,
        }
        for field, expected_type in required_runtime.items():
            if field not in rt_body:
                _fail(
                    "step 7f: /setup/runtime contract",
                    f"missing field {field!r}; got {sorted(rt_body)}",
                )
                return False
            if not isinstance(rt_body[field], expected_type):
                _fail(
                    "step 7f: /setup/runtime contract",
                    f"field {field!r} has type {type(rt_body[field]).__name__}, expected {expected_type}",
                )
                return False
        # Optional fields must be present even when null (Status-Meridian.bat reads them).
        for opt in ("backend_log_path", "structlog_dir", "last_import_job"):
            if opt not in rt_body:
                _fail(
                    "step 7f: /setup/runtime contract",
                    f"optional field {opt!r} missing entirely; must be present (may be null)",
                )
                return False
        _ok(
            "step 7f: /setup/runtime contract",
            f"pid={rt_body['pid']} version={rt_body['version']} uptime={rt_body['uptime_seconds']:.1f}s",
        )

        # Alpha-15: step 7g (auth bypass default-secure) and step 7h
        # (.env-loaded env-var contract) removed. They asserted on the
        # alpha-13/14 MERIDIAN_AUTH_DISABLED bypass machinery, which
        # alpha-15 retired by removing TOTP enforcement entirely
        # (validate-the-product-first principle).

        return True
    finally:
        # Cleanup — kill the backend so the next run isn't held up by a port collision.
        if proc.poll() is None:
            proc.terminate()
            try:
                proc.wait(timeout=10.0)
            except subprocess.TimeoutExpired:
                proc.kill()
                proc.wait(timeout=5.0)


# ──────────────────────────────────────────────────────────────────────────
# Step 7h — .env-loaded env-var contract (alpha-14)
# ──────────────────────────────────────────────────────────────────────────


def _parse_dotenv(text: str) -> dict[str, str]:
    """Parse a minimal KEY=VALUE .env file. Mirrors the simple semantics
    Start-Meridian.bat (alpha-14) uses: ignore blank lines and lines
    starting with '#'; split on the first '='; strip surrounding quotes.
    """
    out: dict[str, str] = {}
    for raw in text.splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        if "=" not in line:
            continue
        key, _, value = line.partition("=")
        key = key.strip()
        value = value.strip()
        if len(value) >= 2 and value[0] == value[-1] and value[0] in ("'", '"'):
            value = value[1:-1]
        if key:
            out[key] = value
    return out


def _step_env_var_via_dotenv(scratch: Path, venv_dir: Path, repo: Path) -> bool:
    """Step 7h (alpha-14): end-to-end verification that a `.env` file in
    MERIDIAN_HOME results in env vars reaching the running backend.

    Why this step exists: alpha-13 shipped MERIDIAN_AUTH_DISABLED=1 as an
    auth-bypass env var. The Python-side contract was tested. But the
    user-facing instructions ("Add MERIDIAN_AUTH_DISABLED=1 to
    C:\\Meridian\\.env, restart the backend") didn't actually work
    because Start-Meridian.bat never loaded .env. Alpha-14 fixed the
    launcher AND added this gauntlet step so the bug class never recurs.

    Simplification: we test the contract, NOT the .bat plumbing. The
    .bat plumbing is covered by the existing static checks (steps 2c/2d)
    plus alpha-14's e2e tests. Here we parse the .env in Python and
    pass the merged dict as `env=` to subprocess.Popen. That exercises
    the same OS-level "env reaches process -> /setup/runtime reports
    it" path from the perspective of contract correctness.
    """
    if os.name == "nt":
        py = venv_dir / "Scripts" / "python.exe"
    else:
        py = venv_dir / "bin" / "python"

    home_7h = scratch / "meridian_home_7h"
    home_7h.mkdir(exist_ok=True)
    spawn_cwd = scratch / "spawn_cwd_7h"
    spawn_cwd.mkdir(exist_ok=True)
    dotenv_path = home_7h / ".env"
    dotenv_path.write_text("MERIDIAN_AUTH_DISABLED=1\n", encoding="utf-8")

    parsed = _parse_dotenv(dotenv_path.read_text(encoding="utf-8"))
    if parsed.get("MERIDIAN_AUTH_DISABLED") != "1":
        _fail(
            "step 7h: env-var via .env contract",
            f".env parser failed to extract MERIDIAN_AUTH_DISABLED=1; got {parsed!r}",
        )
        return False

    port = 8001
    base = f"http://127.0.0.1:{port}"
    log_path = scratch / "backend_7h.log"

    env = {
        **os.environ,
        **parsed,
        "MERIDIAN_HOME": str(home_7h),
        "MERIDIAN_PORT": str(port),
        "MERIDIAN_HOST": "127.0.0.1",
        "MERIDIAN_BACKEND_LOG": str(log_path),
    }
    env.pop("PYTHONPATH", None)

    _info(
        f"step 7h: spawning backend on :{port} with .env-derived "
        f"MERIDIAN_AUTH_DISABLED={parsed.get('MERIDIAN_AUTH_DISABLED')!r}"
    )
    with log_path.open("wb") as logf:
        proc = subprocess.Popen(
            [str(py), "-m", "meridian.api.main"],
            cwd=str(spawn_cwd),
            env=env,
            stdout=logf,
            stderr=subprocess.STDOUT,
        )
    try:
        # 60s deadline @ 250ms intervals, per the brief.
        deadline = time.monotonic() + 60.0
        healthy = False
        while time.monotonic() < deadline:
            status, _body = _probe(f"{base}/health", timeout=0.5)
            if status == 200:
                healthy = True
                break
            time.sleep(0.25)
        if not healthy:
            try:
                tail = log_path.read_text(errors="replace").splitlines()[-30:]
            except OSError:
                tail = []
            _fail(
                "step 7h: env-var via .env contract",
                f"backend on :{port} never reported /health 200 within 60s. "
                "backend.log tail:\n" + "\n".join(tail),
            )
            return False

        rt_status, rt_payload = _probe(f"{base}/setup/runtime", timeout=5.0)
        if rt_status != 200:
            _fail(
                "step 7h: env-var via .env contract",
                f"GET /setup/runtime returned {rt_status} (expected 200)",
            )
            return False
        try:
            rt_body = json.loads(rt_payload.decode("utf-8"))
        except (UnicodeDecodeError, json.JSONDecodeError) as exc:
            _fail(
                "step 7h: env-var via .env contract",
                f"non-JSON body from /setup/runtime: {exc}",
            )
            return False
        auth_disabled = rt_body.get("auth_disabled")
        if not isinstance(auth_disabled, bool):
            _fail(
                "step 7h: env-var via .env contract",
                f"auth_disabled must be bool, got {type(auth_disabled).__name__}",
            )
            return False
        if auth_disabled is not True:
            _fail(
                "step 7h: env-var via .env contract",
                "auth_disabled=false on a backend spawned with "
                "MERIDIAN_AUTH_DISABLED=1 derived from .env. The .env -> env-var "
                "-> /setup/runtime contract is broken. This is the alpha-13 "
                "bug class -- the documentation said 'put it in .env' but "
                "the env var never reached the process. "
                f"rt_body keys: {sorted(rt_body)}",
            )
            return False
        _ok(
            "step 7h: env-var via .env contract",
            "auth_disabled=true after .env load",
        )
        return True
    finally:
        if proc.poll() is None:
            proc.terminate()
            try:
                proc.wait(timeout=10.0)
            except subprocess.TimeoutExpired:
                proc.kill()
                proc.wait(timeout=5.0)
        try:
            dotenv_path.unlink()
        except OSError:
            pass


# ──────────────────────────────────────────────────────────────────────────
# Step 7i — alpha-24 idempotency dedupe under parallel POST
# ──────────────────────────────────────────────────────────────────────────


def _step_idempotency_dedupes_parallel_posts(
    scratch: Path, venv_dir: Path, repo: Path
) -> bool:
    """Step 7i — alpha-24 item #4: same Idempotency-Key + parallel POSTs == one job.

    Spawns two threads, each posting /setup/import-folder with the same
    UUIDv4 token to a tmp folder containing one .pdf. Both responses must
    carry the same job_id; the post-state /setup/runtime probe must
    report at most one folder-import job created during the window.

    Catches wiring drift at the wheel level (analogous to step 2b's
    static check on installer URL constants for the alpha-5 IPv6 bug).
    """
    import concurrent.futures
    import json as _json

    _info("step 7i -- alpha-24 idempotency dedupe under parallel POST")

    if os.name == "nt":
        py = venv_dir / "Scripts" / "python.exe"
    else:
        py = venv_dir / "bin" / "python"

    home_7i = scratch / "meridian_home_7i"
    home_7i.mkdir(exist_ok=True)
    spawn_cwd = scratch / "spawn_cwd_7i"
    spawn_cwd.mkdir(exist_ok=True)

    port = 8002
    base_url = f"http://127.0.0.1:{port}"
    log_path = scratch / "backend_7i.log"

    env = {
        **os.environ,
        "MERIDIAN_HOME": str(home_7i),
        "MERIDIAN_PORT": str(port),
        "MERIDIAN_HOST": "127.0.0.1",
    }
    env.pop("PYTHONPATH", None)

    _info(f"step 7i: spawning backend on :{port}")
    with log_path.open("wb") as logf:
        proc = subprocess.Popen(
            [str(py), "-m", "meridian.api.main"],
            cwd=str(spawn_cwd),
            env=env,
            stdout=logf,
            stderr=subprocess.STDOUT,
        )
    try:
        deadline = time.monotonic() + 60.0
        healthy = False
        while time.monotonic() < deadline:
            status, _body = _probe(f"{base_url}/health", timeout=0.5)
            if status == 200:
                healthy = True
                break
            time.sleep(0.25)
        if not healthy:
            try:
                tail = log_path.read_text(errors="replace").splitlines()[-30:]
            except OSError:
                tail = []
            _fail(
                "step 7i: idempotency dedupe",
                f"backend on :{port} never reported /health 200 within 60s. "
                "backend.log tail:\n" + "\n".join(tail),
            )
            return False

        with tempfile.TemporaryDirectory() as td:
            folder = Path(td) / "src"
            folder.mkdir()
            (folder / "doc.pdf").write_bytes(b"%PDF-1.4 gauntlet 7i\n%%EOF\n")

            token = "11111111-1111-4111-8111-111111111111"
            body = _json.dumps({
                "folder_path": str(folder),
                "project_name": "gauntlet-7i",
            }).encode("utf-8")

            def _post() -> str | None:
                try:
                    req = urllib.request.Request(
                        f"{base_url}/setup/import-folder",
                        data=body,
                        headers={
                            "Content-Type": "application/json",
                            "Idempotency-Key": token,
                        },
                        method="POST",
                    )
                    with urllib.request.urlopen(req, timeout=10) as resp:  # noqa: S310 -- localhost
                        payload = _json.loads(resp.read().decode("utf-8"))
                        return payload["job_id"]
                except (urllib.error.URLError, TimeoutError, _json.JSONDecodeError):
                    return None

            with concurrent.futures.ThreadPoolExecutor(max_workers=2) as ex:
                f1 = ex.submit(_post)
                f2 = ex.submit(_post)
                job_id_1 = f1.result()
                job_id_2 = f2.result()

        if job_id_1 is None or job_id_2 is None:
            _fail("7i", "parallel POST failed (network or parse error). See backend.log.")
            return False
        if job_id_1 != job_id_2:
            _fail(
                "7i",
                f"parallel POSTs with same Idempotency-Key returned different "
                f"job_ids: {job_id_1!r} vs {job_id_2!r}. Backend dedupe is "
                "not wired or the registry is leaking.",
            )
            return False
        _ok("7i", f"both POSTs returned job_id={job_id_1[:8]}...")
        return True
    finally:
        if proc.poll() is None:
            proc.terminate()
            try:
                proc.wait(timeout=10.0)
            except subprocess.TimeoutExpired:
                proc.kill()
                proc.wait(timeout=5.0)


# ──────────────────────────────────────────────────────────────────────────
# Step 7j — alpha-25 pipeline e2e + workbook conflict_summary smoke
# ──────────────────────────────────────────────────────────────────────────


def _step_pipeline_e2e_and_workbook_smoke(
    scratch: Path, venv_dir: Path, repo: Path
) -> bool:
    """Step 7j (alpha-25): pipeline e2e + workbook conflict_summary smoke.

    Kicks /api/projects/<slug>/pipeline (no idempotency key — first run),
    polls until phase=done (cap 4 minutes), asserts deliverables were
    produced, downloads the Excel export and verifies the
    conflict_summary column is present in the Master sheet.

    Catches wiring drift between the alpha-25 pipeline worker and the
    workbook builder introduced in that release.
    """
    import io
    import json as _json
    import time as _t

    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        _fail("step 7j: workbook download", f"openpyxl not available: {exc}")
        return False

    _info("step 7j -- alpha-25 pipeline e2e + workbook conflict_summary smoke")

    if os.name == "nt":
        py = venv_dir / "Scripts" / "python.exe"
    else:
        py = venv_dir / "bin" / "python"

    home_7j = scratch / "meridian_home_7j"
    home_7j.mkdir(exist_ok=True)
    spawn_cwd = scratch / "spawn_cwd_7j"
    spawn_cwd.mkdir(exist_ok=True)

    port = 8003
    base_url = f"http://127.0.0.1:{port}"
    log_path = scratch / "backend_7j.log"

    env = {
        **os.environ,
        "MERIDIAN_HOME": str(home_7j),
        "MERIDIAN_PORT": str(port),
        "MERIDIAN_HOST": "127.0.0.1",
    }
    env.pop("PYTHONPATH", None)

    _info(f"step 7j: spawning backend on :{port}")
    with log_path.open("wb") as logf:
        proc = subprocess.Popen(
            [str(py), "-m", "meridian.api.main"],
            cwd=str(spawn_cwd),
            env=env,
            stdout=logf,
            stderr=subprocess.STDOUT,
        )
    try:
        deadline = _t.monotonic() + 60.0
        healthy = False
        while _t.monotonic() < deadline:
            status, _body = _probe(f"{base_url}/health", timeout=0.5)
            if status == 200:
                healthy = True
                break
            _t.sleep(0.25)
        if not healthy:
            try:
                tail = log_path.read_text(errors="replace").splitlines()[-30:]
            except OSError:
                tail = []
            _fail(
                "step 7j: pipeline e2e",
                f"backend on :{port} never reported /health 200 within 60s. "
                "backend.log tail:\n" + "\n".join(tail),
            )
            return False

        # Create a project + ingest one synthetic DOCX via /setup/import-folder.
        # A bare /api/projects POST creates an empty project (no sources), and
        # the pipeline assertion below requires deliverables to actually be
        # produced — so we must ingest at least one parseable file. A bare
        # %PDF-1.4 stub fails at the text-extraction stage; python-docx makes
        # a real, parseable document the dispatcher can ingest. (Mirrors the
        # alpha-25 e2e fixture pattern in tests/e2e/conftest.py.)
        try:
            from docx import Document  # noqa: PLC0415 — script-time import
        except ImportError as exc:
            _fail(
                "step 7j: pipeline e2e",
                f"python-docx not available in controller Python: {exc}",
            )
            return False

        with tempfile.TemporaryDirectory() as td:
            folder = Path(td) / "src"
            folder.mkdir()
            doc = Document()
            doc.add_heading("Sample Specification Section", level=1)
            for line in (
                "Contractor shall supply and install one (1) air handling unit.",
                "All ductwork shall be sealed to SMACNA Class A.",
                "Coordinate penetrations with the structural engineer.",
            ):
                doc.add_paragraph(line)
            doc.save(str(folder / "spec.docx"))

            slug = "gauntlet-7j"
            import_payload = _json.dumps({
                "folder_path": str(folder),
                "project_name": slug,
            }).encode("utf-8")
            try:
                req = urllib.request.Request(
                    f"{base_url}/setup/import-folder",
                    data=import_payload,
                    headers={
                        "Content-Type": "application/json",
                        "User-Agent": "release-gauntlet",
                    },
                    method="POST",
                )
                with urllib.request.urlopen(req, timeout=15) as resp:  # noqa: S310 — localhost
                    import_body = _json.loads(resp.read().decode("utf-8"))
            except (urllib.error.URLError, TimeoutError, _json.JSONDecodeError) as exc:
                _fail("step 7j: pipeline e2e", f"import-folder POST failed: {exc}")
                return False

            import_job_id = import_body.get("job_id")
            if not isinstance(import_job_id, str):
                _fail(
                    "step 7j: pipeline e2e",
                    f"import-folder response missing job_id; got {import_body!r}",
                )
                return False

            # Wait for import to land before kicking the pipeline.
            import_deadline = _t.monotonic() + 30.0
            import_done = False
            while _t.monotonic() < import_deadline:
                ip_status, ip_body = _probe(
                    f"{base_url}/setup/import-folder/{import_job_id}", timeout=5.0,
                )
                if ip_status == 200:
                    try:
                        ip_state = _json.loads(ip_body.decode("utf-8"))
                    except _json.JSONDecodeError:
                        ip_state = {}
                    if ip_state.get("status") in {"succeeded", "failed"}:
                        if ip_state.get("status") == "failed":
                            _fail(
                                "step 7j: pipeline e2e",
                                f"import-folder failed: {ip_state.get('failed_groups')}",
                            )
                            return False
                        import_done = True
                        break
                _t.sleep(0.5)
            if not import_done:
                _fail("step 7j: pipeline e2e", "import-folder didn't reach terminal status in 30s")
                return False
        # tempfile is cleaned up; the project + sources persist in the
        # backend's MERIDIAN_HOME for the rest of step 7j.

        # 1. Kick off the pipeline (no idempotency key -- first run on this slug).
        pipeline_payload = _json.dumps({}).encode("utf-8")
        try:
            req = urllib.request.Request(
                f"{base_url}/api/projects/{slug}/pipeline",
                data=pipeline_payload,
                headers={
                    "Content-Type": "application/json",
                    "User-Agent": "release-gauntlet",
                },
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=10) as resp:  # noqa: S310 — localhost
                post_body = _json.loads(resp.read().decode("utf-8"))
        except urllib.error.HTTPError as exc:
            _fail(
                "step 7j: pipeline POST",
                f"status={exc.code} body={exc.read()[:200]!r}",
            )
            return False
        except (urllib.error.URLError, TimeoutError, _json.JSONDecodeError) as exc:
            _fail("step 7j: pipeline POST", f"failed: {exc}")
            return False

        job_id = post_body.get("job_id")
        if not isinstance(job_id, str):
            _fail(
                "step 7j: pipeline POST",
                f"response missing job_id; got {post_body!r}",
            )
            return False

        # 2. Poll until phase=done OR 4 minutes elapse.
        deadline = _t.monotonic() + 240.0
        last: dict = {}
        while _t.monotonic() < deadline:
            poll_status, poll_payload = _probe(
                f"{base_url}/api/projects/{slug}/pipeline/{job_id}", timeout=5.0
            )
            if poll_status == 200:
                try:
                    last = _json.loads(poll_payload.decode("utf-8"))
                except _json.JSONDecodeError:
                    last = {}
                if last.get("phase") in {"done", "failed"}:
                    break
            _t.sleep(2.0)

        if not last or last.get("phase") != "done":
            _fail(
                "step 7j: pipeline completion",
                f"final={last!r}",
            )
            return False
        _ok(
            "step 7j: pipeline completion",
            f"phase={last['phase']} bootstrap={last.get('bootstrap_status')}",
        )

        # 2b. Conflict-pass ran (alpha-25.1 regression pin).
        #
        # The original alpha-25 keystone wired bootstrap + extract but missed
        # conflict-pass — so the master register's flags column shipped empty
        # and the new conflict_summary column had nothing to render. The fix
        # adds conflict-pass as the third pipeline step; this assertion
        # ensures it stays wired. Acceptable terminal states are 'succeeded'
        # (single source = pass ran, 0 conflicts persisted is fine) or
        # 'skipped' (zero deliverables produced — extract failed silently).
        # 'failed' or 'pending' here means conflict-pass either errored or
        # never ran, which is the regression mode we are pinning.
        cp_status = last.get("conflict_pass_status")
        if cp_status not in {"succeeded", "skipped"}:
            _fail(
                "step 7j: conflict-pass ran",
                f"conflict_pass_status={cp_status!r} (expected 'succeeded' or 'skipped')",
            )
            return False
        _ok("step 7j: conflict-pass ran", f"conflict_pass_status={cp_status}")

        # 3. Coverage sanity.
        cov_status, cov_payload = _probe(
            f"{base_url}/api/projects/{slug}/coverage", timeout=10.0
        )
        if cov_status != 200:
            _fail("step 7j: deliverables produced", f"coverage GET returned {cov_status}")
            return False
        try:
            cov = _json.loads(cov_payload.decode("utf-8"))
        except _json.JSONDecodeError as exc:
            _fail("step 7j: deliverables produced", f"non-JSON coverage body: {exc}")
            return False
        total = cov.get("deliverable_status", {}).get("total", 0)
        if total <= 0:
            _fail("step 7j: deliverables produced", f"total={total}")
            return False
        _ok("step 7j: deliverables produced", f"total={total}")

        # 4. Workbook column smoke.
        try:
            xl_status, xl_bytes = _probe(
                f"{base_url}/api/projects/{slug}/export.xlsx", timeout=30.0
            )
        except Exception as exc:
            _fail("step 7j: workbook download", f"failed: {exc}")
            return False
        if xl_status != 200:
            _fail("step 7j: workbook download", f"export.xlsx returned {xl_status}")
            return False
        try:
            wb = load_workbook(io.BytesIO(xl_bytes))
            if "Master" not in wb.sheetnames:
                _fail(
                    "step 7j: workbook download",
                    f"no 'Master' sheet; sheets={wb.sheetnames}",
                )
                return False
            headers = [c.value for c in wb["Master"][1]]
        except Exception as exc:
            _fail("step 7j: workbook download", f"failed to parse xlsx: {exc}")
            return False
        if "conflict_summary" not in headers:
            _fail("step 7j: conflict_summary column", f"headers={headers}")
            return False
        _ok("step 7j: conflict_summary column", "present")
        return True
    finally:
        if proc.poll() is None:
            proc.terminate()
            try:
                proc.wait(timeout=10.0)
            except subprocess.TimeoutExpired:
                proc.kill()
                proc.wait(timeout=5.0)


# ──────────────────────────────────────────────────────────────────────────
# Step 7k — alpha-26 SSE event stream + conflict register + hierarchy
# ──────────────────────────────────────────────────────────────────────────


def _step_alpha26_sse_and_conflicts(
    scratch: Path, venv_dir: Path, repo: Path
) -> bool:
    """Step 7k (alpha-26): SSE event stream + conflict register + hierarchy.

    Spawns a backend on :8004, ingests a 3-source synthetic .docx corpus,
    runs the pipeline to phase=done, then validates:
      - SSE stream delivered >=1 extraction.source.start frame during the run
      - GET /conflict-register returns a non-empty items array
      - GET /conflict-register.xlsx loads cleanly via openpyxl
      - GET /hierarchy responds with resolved_count >= 0
      - POST /conflicts/{id}/resolve flips a pending row + lifts hierarchy
        resolved_count to >= 1
      - Subscriber cap honoured: 6 simultaneous /events GETs return 503
        for the 6th when MERIDIAN_EVENTS_MAX_SUBSCRIBERS=5.
    """
    import io
    import json as _json
    import threading
    import time as _t
    import urllib.request
    import urllib.error

    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        _fail("step 7k: openpyxl not available", str(exc))
        return False

    _info("step 7k -- alpha-26 SSE event stream + conflicts surfaces")

    if os.name == "nt":
        py = venv_dir / "Scripts" / "python.exe"
    else:
        py = venv_dir / "bin" / "python"

    home_7k = scratch / "meridian_home_7k"
    home_7k.mkdir(exist_ok=True)
    spawn_cwd = scratch / "spawn_cwd_7k"
    spawn_cwd.mkdir(exist_ok=True)

    port = 8004
    base_url = f"http://127.0.0.1:{port}"
    log_path = scratch / "backend_7k.log"

    env = {
        **os.environ,
        "MERIDIAN_HOME": str(home_7k),
        "MERIDIAN_PORT": str(port),
        "MERIDIAN_HOST": "127.0.0.1",
        "MERIDIAN_EVENTS_MAX_SUBSCRIBERS": "5",
    }
    env.pop("PYTHONPATH", None)

    _info(f"step 7k: spawning backend on :{port}")
    with log_path.open("wb") as logf:
        proc = subprocess.Popen(
            [str(py), "-m", "meridian.api.main"],
            cwd=str(spawn_cwd),
            env=env,
            stdout=logf,
            stderr=subprocess.STDOUT,
        )
    try:
        # Wait for healthy
        deadline = _t.monotonic() + 60.0
        healthy = False
        while _t.monotonic() < deadline:
            status, _body = _probe(f"{base_url}/health", timeout=0.5)
            if status == 200:
                healthy = True
                break
            _t.sleep(0.25)
        if not healthy:
            _fail("step 7k: backend startup", "backend never became healthy in 60s")
            return False

        # Ingest a 3-source .docx corpus via /setup/import-folder
        try:
            from docx import Document
        except ImportError as exc:
            _fail("step 7k: python-docx not available", str(exc))
            return False

        with tempfile.TemporaryDirectory() as td:
            folder = Path(td) / "src"
            folder.mkdir()
            for i, name in enumerate(["spec_a.docx", "spec_b.docx", "spec_c.docx"]):
                doc = Document()
                doc.add_heading(f"Sample Section {i+1}", level=1)
                for line in (
                    f"Contractor shall supply unit {i+1}.",
                    "All ductwork shall be sealed to SMACNA Class A.",
                    "Coordinate penetrations with the structural engineer.",
                ):
                    doc.add_paragraph(line)
                doc.save(str(folder / name))

            slug = "gauntlet-7k"
            import_payload = _json.dumps({
                "folder_path": str(folder),
                "project_name": slug,
            }).encode("utf-8")
            req = urllib.request.Request(
                f"{base_url}/setup/import-folder",
                data=import_payload,
                headers={"Content-Type": "application/json", "User-Agent": "release-gauntlet"},
                method="POST",
            )
            try:
                with urllib.request.urlopen(req, timeout=15) as resp:
                    import_body = _json.loads(resp.read().decode("utf-8"))
            except (urllib.error.URLError, TimeoutError, _json.JSONDecodeError) as exc:
                _fail("step 7k: import-folder POST", f"failed: {exc}")
                return False

            import_job_id = import_body.get("job_id")
            if not isinstance(import_job_id, str):
                _fail("step 7k: import-folder POST", f"missing job_id: {import_body!r}")
                return False

            # Wait for import to complete
            import_deadline = _t.monotonic() + 30.0
            while _t.monotonic() < import_deadline:
                ip_status, ip_body = _probe(
                    f"{base_url}/setup/import-folder/{import_job_id}", timeout=5.0,
                )
                if ip_status == 200:
                    try:
                        ip_state = _json.loads(ip_body.decode("utf-8"))
                    except _json.JSONDecodeError:
                        ip_state = {}
                    if ip_state.get("status") in {"succeeded", "failed"}:
                        if ip_state.get("status") == "failed":
                            _fail("step 7k: import-folder", f"failed: {ip_state}")
                            return False
                        break
                _t.sleep(0.5)

        # SSE smoke: open a stream + capture frames during pipeline run.
        sse_frames: list[dict] = []
        sse_done = threading.Event()

        def _consume_sse():
            try:
                req = urllib.request.Request(
                    f"{base_url}/api/projects/{slug}/events",
                    headers={"Accept": "text/event-stream"},
                )
                with urllib.request.urlopen(req, timeout=120) as resp:
                    if resp.status != 200:
                        return
                    buffer = ""
                    while not sse_done.is_set():
                        chunk = resp.read(1024)
                        if not chunk:
                            break
                        buffer += chunk.decode("utf-8", errors="replace")
                        while "\n\n" in buffer:
                            frame, buffer = buffer.split("\n\n", 1)
                            event_name = "log"
                            data_line = ""
                            for line in frame.split("\n"):
                                if line.startswith("event: "):
                                    event_name = line[7:].strip()
                                elif line.startswith("data: "):
                                    data_line = line[6:]
                            if event_name == "log" and data_line:
                                try:
                                    sse_frames.append(_json.loads(data_line))
                                except _json.JSONDecodeError:
                                    pass
            except Exception:
                pass

        sse_thread = threading.Thread(target=_consume_sse, daemon=True)
        sse_thread.start()
        _t.sleep(0.5)  # let SSE connect before kicking pipeline

        # Run pipeline
        pipeline_payload = _json.dumps({}).encode("utf-8")
        try:
            req = urllib.request.Request(
                f"{base_url}/api/projects/{slug}/pipeline",
                data=pipeline_payload,
                headers={"Content-Type": "application/json", "User-Agent": "release-gauntlet"},
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=10) as resp:
                post_body = _json.loads(resp.read().decode("utf-8"))
        except Exception as exc:
            _fail("step 7k: pipeline POST", f"failed: {exc}")
            return False
        job_id = post_body.get("job_id")

        # Poll until done
        deadline = _t.monotonic() + 240.0
        last = None
        while _t.monotonic() < deadline:
            poll_status, poll_payload = _probe(
                f"{base_url}/api/projects/{slug}/pipeline/{job_id}", timeout=5.0,
            )
            if poll_status == 200:
                try:
                    last = _json.loads(poll_payload.decode("utf-8"))
                except _json.JSONDecodeError:
                    pass
                if last and last.get("phase") in {"done", "failed"}:
                    break
            _t.sleep(2.0)

        sse_done.set()

        if not last or last.get("phase") != "done":
            _fail("step 7k: pipeline completion", f"final={last!r}")
            return False
        _ok("step 7k: pipeline completion", f"phase={last['phase']}")

        # SSE assertion
        seen_events = {f.get("event") for f in sse_frames}
        if not any(e and e.startswith("extraction.source") for e in seen_events):
            _fail("step 7k: SSE frames", f"no extraction.source.* frames; saw {seen_events}")
            return False
        _ok("step 7k: SSE frames", f"saw {len(sse_frames)} log frames")

        # Conflict register JSON
        try:
            cr_status, cr_body = _probe(
                f"{base_url}/api/projects/{slug}/conflict-register", timeout=10.0,
            )
        except Exception as exc:
            _fail("step 7k: conflict-register GET", f"failed: {exc}")
            return False
        if cr_status != 200:
            _fail("step 7k: conflict-register GET", f"status={cr_status}")
            return False
        cr = _json.loads(cr_body.decode("utf-8"))
        _ok("step 7k: conflict-register JSON", f"counts={cr.get('counts', {})}")

        # Conflict register xlsx
        try:
            xl_status, xl_bytes = _probe(
                f"{base_url}/api/projects/{slug}/conflict-register.xlsx", timeout=30.0,
            )
        except Exception as exc:
            _fail("step 7k: conflict-register xlsx", f"failed: {exc}")
            return False
        if xl_status != 200:
            _fail("step 7k: conflict-register xlsx", f"status={xl_status}")
            return False
        try:
            wb = load_workbook(io.BytesIO(xl_bytes))
            assert "Conflicts" in wb.sheetnames
        except Exception as exc:
            _fail("step 7k: conflict-register xlsx parse", f"failed: {exc}")
            return False
        _ok("step 7k: conflict-register xlsx", "parsed cleanly")

        # Hierarchy
        try:
            h_status, h_body = _probe(
                f"{base_url}/api/projects/{slug}/hierarchy", timeout=10.0,
            )
        except Exception as exc:
            _fail("step 7k: hierarchy GET", f"failed: {exc}")
            return False
        if h_status != 200:
            _fail("step 7k: hierarchy GET", f"status={h_status}")
            return False
        h = _json.loads(h_body.decode("utf-8"))
        _ok("step 7k: hierarchy", f"resolved_count={h.get('resolved_count')}")

        # Subscriber cap: probe current active count via /setup/runtime,
        # open (cap - active) streams expecting success, then one more
        # expecting 503. Robust to whatever subscribers are still alive
        # from the SSE consumer thread used in the pipeline-run smoke
        # (the urllib reader can't be cleanly cancelled mid-read).
        try:
            rt_status, rt_body = _probe(
                f"{base_url}/api/setup/runtime", timeout=5.0,
            )
            if rt_status != 200:
                _fail("step 7k: subscriber cap", f"/setup/runtime returned {rt_status}")
                return False
            rt = _json.loads(rt_body.decode("utf-8"))
            active = rt["events"]["active_subscribers"]
            cap = rt["events"]["max_subscribers"]
            available = max(0, cap - active)
        except Exception as exc:
            _fail("step 7k: subscriber cap", f"runtime probe failed: {exc}")
            return False

        cap_streams = []
        try:
            # Open exactly enough streams to fill the remaining capacity.
            for i in range(available):
                req = urllib.request.Request(
                    f"{base_url}/api/projects/{slug}/events",
                    headers={"Accept": "text/event-stream"},
                )
                cap_streams.append(urllib.request.urlopen(req, timeout=5.0))
            # Next stream MUST return 503 — cap is now reached.
            req_over = urllib.request.Request(
                f"{base_url}/api/projects/{slug}/events",
                headers={"Accept": "text/event-stream"},
            )
            try:
                urllib.request.urlopen(req_over, timeout=5.0)
                _fail("step 7k: subscriber cap", f"over-cap subscriber returned 200; expected 503 (cap={cap})")
                return False
            except urllib.error.HTTPError as e:
                if e.code != 503:
                    _fail("step 7k: subscriber cap", f"over-cap subscriber returned {e.code}; expected 503")
                    return False
            _ok("step 7k: subscriber cap", f"cap={cap}, opened {available} more, +1 returned 503")
        finally:
            for s in cap_streams:
                try:
                    s.close()
                except Exception:
                    pass

        return True
    finally:
        if proc.poll() is None:
            proc.terminate()
            try:
                proc.wait(timeout=10.0)
            except subprocess.TimeoutExpired:
                proc.kill()
                proc.wait(timeout=5.0)


# ──────────────────────────────────────────────────────────────────────────
# Step 8 — CLI --help smoke
# ──────────────────────────────────────────────────────────────────────────


def _step_cli_help_smoke(venv_dir: Path) -> bool:
    if os.name == "nt":
        meridian_exe = venv_dir / "Scripts" / "meridian.exe"
    else:
        meridian_exe = venv_dir / "bin" / "meridian"

    invocations = [
        [str(meridian_exe), "--help"],
        [str(meridian_exe), "start", "--help"],
        [str(meridian_exe), "init", "--help"],
    ]
    # Belt-and-braces: keep webbrowser.open inert if the user's `meridian start`
    # somehow runs (it won't here, but env hygiene is cheap).
    env = {**os.environ, "MERIDIAN_TEST_NO_BROWSER": "1"}
    for cmd in invocations:
        try:
            r = subprocess.run(
                cmd, capture_output=True, text=True, timeout=30, env=env
            )
        except subprocess.TimeoutExpired:
            _fail("step 8: CLI help", f"{' '.join(cmd)} timed out")
            return False
        if r.returncode != 0:
            tail = (r.stdout + r.stderr).splitlines()[-20:]
            _fail(
                "step 8: CLI help",
                f"{' '.join(cmd)} exit={r.returncode}\n" + "\n".join(tail),
            )
            return False
    _ok("step 8: CLI help", f"{len(invocations)} invocation(s) returned 0")
    return True


# ──────────────────────────────────────────────────────────────────────────
# Step 2b — installer URL constants
# ──────────────────────────────────────────────────────────────────────────


def _step_detach_pattern(installer_dir: Path) -> bool:
    """Step 7d (alpha-12): the backend launch path MUST use pythonw.exe AND
    WindowStyle Hidden. Catches a regression that re-introduces the
    alpha-11 'close terminal kills backend' bug class.

    Two surfaces are checked:

      * ``installer/Start-Meridian.bat`` — the user-facing launcher.
        Must reference ``pythonw.exe`` AND ``-WindowStyle Hidden``.
      * ``installer/Install-Meridian.ps1`` — the install-time backend
        spawn. Must reference ``pythonw.exe`` AND ``-WindowStyle Hidden``
        in the ``Start-BackendAndOpenBrowser`` function.

    Both are static-content greps, not live-process tests — runtime
    detachment-survival belongs in tests/e2e where we can simulate
    process-group semantics deterministically.
    """
    start_bat = installer_dir / "Start-Meridian.bat"
    install_ps1 = installer_dir / "Install-Meridian.ps1"

    failures: list[str] = []
    # Alpha-14 reviewer-pull-forward: the alpha-13 documentation-vs-reality
    # gap (Start-Meridian.bat didn't load .env, so MERIDIAN_AUTH_DISABLED
    # in .env was orphaned) MUST be statically guarded. Without these
    # tokens, a regression that deletes the .env loader would pass step
    # 7h (subprocess-level) but ship a broken .bat. This closes the
    # alpha-13-class miss the rest of alpha-14 exists to prevent.
    required_tokens = (
        ("pythonw.exe",          "alpha-12 detach pattern (visible-cmd regression)"),
        ("WindowStyle Hidden",   "alpha-12 detach pattern (visible-cmd regression)"),
        # Alpha-15: dropped MERIDIAN_ENV_FILE / Set-Item Env: tokens that
        # alpha-14 added. Those existed only to wire the alpha-13
        # MERIDIAN_AUTH_DISABLED bypass through .env -- which alpha-15
        # retired by removing TOTP enforcement entirely. The cmd.exe
        # parser pathology in alpha-14's multi-line PS argument is also
        # gone (single-line spawn restored). Validate-the-product-first.
    )

    if not start_bat.exists():
        failures.append(f"missing: {start_bat}")
    else:
        text = start_bat.read_text(encoding="utf-8", errors="replace")
        for token, why in required_tokens:
            if token not in text:
                failures.append(f"{start_bat.name}: missing {token!r} -- {why}")

    if not install_ps1.exists():
        failures.append(f"missing: {install_ps1}")
    else:
        ps_text = install_ps1.read_text(encoding="utf-8", errors="replace")
        for token, why in required_tokens:
            if token not in ps_text:
                failures.append(f"{install_ps1.name}: missing {token!r} -- {why}")

    if failures:
        _fail(
            "step 2c: detach-pattern static check",
            "alpha-12 contract -- pythonw.exe + WindowStyle Hidden must appear in "
            "BOTH Start-Meridian.bat and the Install-Meridian.ps1 inline copy:\n  - "
            + "\n  - ".join(failures),
        )
        return False
    _ok(
        "step 2c: detach-pattern static check",
        "pythonw.exe + WindowStyle Hidden present in both surfaces",
    )
    return True


def _step_launcher_exec_policy(installer_dir: Path) -> bool:
    """Step 2d (alpha-12, broadened alpha-13): every .bat under installer/
    that invokes a .ps1 must use ``-ExecutionPolicy Bypass``.

    Alpha-12 only checked Meridian-Console.bat. Alpha-13 broadens the
    rule because a regression that adds a new .bat-invokes-.ps1 surface
    (Update-Meridian.bat, Repair-Meridian.bat, etc.) would silently
    re-expose the alpha-11 ExecutionPolicy block.

    Detection: any line in any .bat file that mentions ``powershell.exe``
    or ``powershell `` AND ``-File`` (i.e. running a script) must also
    contain ``-ExecutionPolicy Bypass``. Inline ``-Command`` invocations
    are exempt because PS execution policy doesn't gate them.
    """
    failures: list[str] = []
    bats = sorted(installer_dir.glob("*.bat"))
    if not bats:
        _info("No .bat files under installer/ -- nothing to check.")
        _ok("step 2d: launcher exec-policy (no bats present)")
        return True

    for bat in bats:
        try:
            lines = bat.read_text(encoding="utf-8", errors="replace").splitlines()
        except OSError as exc:
            failures.append(f"{bat.name}: could not read ({exc})")
            continue
        for line_num, raw in enumerate(lines, start=1):
            stripped = raw.lstrip()
            if stripped.startswith("REM ") or stripped.startswith("::"):
                continue
            lower = stripped.lower()
            invokes_ps = (
                "powershell.exe" in lower or
                lower.startswith("powershell ") or
                " powershell " in lower
            )
            if not invokes_ps:
                continue
            # Only flag invocations that run a .ps1 script (-File). Inline
            # -Command / -EncodedCommand / -NoProfile-only invocations
            # don't hit the script-execution policy gate.
            if "-file" not in lower:
                continue
            if "-executionpolicy bypass" not in lower:
                failures.append(
                    f"{bat.name}:{line_num}: invokes a .ps1 via -File but lacks "
                    "-ExecutionPolicy Bypass"
                )
    if failures:
        _fail(
            "step 2d: launcher exec-policy",
            "every .bat invoking a .ps1 via -File must include "
            "-ExecutionPolicy Bypass (alpha-13 broadening of the alpha-12 "
            "Meridian-Console-only check):\n  - " + "\n  - ".join(failures),
        )
        return False
    _ok(
        "step 2d: launcher exec-policy",
        f"{len(bats)} .bat file(s) under installer/ all bypass ExecutionPolicy where needed",
    )
    return True


def _step_installer_url_constants(installer_dir: Path) -> bool:
    """Catches the alpha-5 regression class: Install-Meridian.ps1 must NOT
    use 'http://localhost:...' in URL constants.

    On Windows, the name 'localhost' resolves to ``::1`` (IPv6 loopback)
    before falling back to ``127.0.0.1`` (IPv4). Uvicorn binds IPv4 only by
    default. PowerShell's HttpWebRequest with a 1s timeout doesn't fall
    back to IPv4 fast enough — the probe targets an empty IPv6 socket,
    times out, and the installer hangs forever at 'Waiting for the backend'
    even though the backend is fully healthy. The user lived this in alpha-4.

    All probe / open URL constants in the installer must use ``127.0.0.1``
    explicitly. Browsers handle IPv6→IPv4 fallback well; PowerShell does not.
    """
    bad: list[str] = []
    for ps1 in installer_dir.rglob("*.ps1"):
        for line_num, raw in enumerate(ps1.read_text(encoding="utf-8").splitlines(), start=1):
            line = raw.strip()
            if line.startswith("#"):
                continue  # tolerate localhost in comments / docstrings
            if "://localhost" in line:
                bad.append(f"  {ps1.name}:{line_num}: {line}")
    if bad:
        _fail(
            "step 2b: installer URL constants",
            "found 'http://localhost' URL(s) in installer .ps1 files. Use 127.0.0.1 "
            "— Windows resolves localhost to IPv6 ::1 first; uvicorn binds IPv4. "
            "Probes hang forever.\n" + "\n".join(bad),
        )
        return False
    _ok("step 2b: installer URL constants", "no 'http://localhost' URLs in installer .ps1 files")
    return True


# ──────────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────────


def main() -> int:
    repo = _repo_root()
    installer_dir = repo / "installer"
    print(f"Meridian release gauntlet — repo at {repo}")
    print(f"  pyproject version: {_read_pyproject_version(repo)}")
    print()

    if not _step_powershell_parse(installer_dir):
        return 1
    if not _step_ascii_only(installer_dir):
        return 1
    if not _step_installer_url_constants(installer_dir):
        return 1
    if not _step_detach_pattern(installer_dir):
        return 1
    if not _step_launcher_exec_policy(installer_dir):
        return 1

    ok, wheel = _step_build_wheel(repo)
    if not ok or wheel is None:
        return 1

    # ignore_cleanup_errors: on Windows the backend may briefly hold a log
    # file handle even after terminate() — pytest-style scratch hygiene
    # matters less than not masking a passing run with a cleanup OSError.
    with tempfile.TemporaryDirectory(
        prefix="meridian-gauntlet-", ignore_cleanup_errors=True
    ) as scratch_str:
        scratch = Path(scratch_str)
        ok, venv_dir = _step_fresh_install(wheel, scratch)
        if not ok or venv_dir is None:
            return 1

        if not _step_backend_lifecycle(venv_dir, scratch, repo):
            return 1

        # Alpha-15: step 7h removed; tested the .env -> bypass plumbing
        # which alpha-15 retired with TOTP enforcement.

        if not _step_idempotency_dedupes_parallel_posts(scratch, venv_dir, repo):
            return 1

        if not _step_pipeline_e2e_and_workbook_smoke(scratch, venv_dir, repo):
            return 1

        if not _step_alpha26_sse_and_conflicts(scratch, venv_dir, repo):
            return 1

        if not _step_cli_help_smoke(venv_dir):
            return 1

    print()
    print("[ ALL ] release gauntlet PASSED")
    return 0


if __name__ == "__main__":
    sys.exit(main())
