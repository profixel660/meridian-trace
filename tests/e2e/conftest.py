"""Shared fixtures for the Meridian end-to-end test harness.

Every fixture is offline-safe: the LLM client is monkeypatched so no network
call ever happens, and project SQLite files live under pytest's tmp_path so
each test runs in an isolated workspace.

Determinism notes:
- ``settings.data_dir`` is repointed at a per-test tmp dir before each test
  imports anything project-related, then restored.
- The ``mock_llm_client`` fixture replaces ``meridian.llm.client.call_llm``
  with a stub that returns canned ``LlmCall`` records and persists them to
  the ``llm_call`` table the same way the real client does.
"""

from __future__ import annotations

import contextlib
import json
import sqlite3
import uuid
from collections.abc import Callable, Iterator
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from fastapi.testclient import TestClient

import time

from meridian import llm as _llm_pkg  # noqa: F401  (ensure package import)
from meridian.config import settings
from meridian.db.connection import connect, transaction
from meridian.llm.client import LlmCall
from meridian.projects import create_project

# --------------------------------------------------------------------------
# Projects-dir isolation
# --------------------------------------------------------------------------


@pytest.fixture
def tmp_projects_dir(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    """Repoint ``settings.data_dir`` at a per-test tmp directory.

    The Meridian config exposes ``settings.projects_dir`` as a derived
    property of ``settings.data_dir``; overriding the latter in-place gives
    us full filesystem isolation without re-importing the settings module.

    Alpha-22: also redirect the wizard state file via ``MERIDIAN_WIZARD_STATE_DIR``
    so the new stable ``state_path()`` (which resolves to ``~/.meridian/``
    rather than ``<projects_dir>/_meridian/``) stays isolated per test.
    Without this, tests share the developer's real ``~/.meridian/onboarding_state.json``
    and get contaminated state.
    """
    target = tmp_path / "projects"
    target.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(settings, "data_dir", target)
    monkeypatch.setenv("MERIDIAN_DATA_DIR", str(target))
    # Redirect the wizard state file to a per-test directory so tests do
    # not share (or corrupt) the developer's real ~/.meridian/ state.
    wizard_state_dir = tmp_path / "wizard_state"
    wizard_state_dir.mkdir(parents=True, exist_ok=True)
    monkeypatch.setenv("MERIDIAN_WIZARD_STATE_DIR", str(wizard_state_dir))
    # Belt-and-braces: clear any cloud creds so any accidental real call
    # would fail loudly rather than silently spend money.
    monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    return target


# --------------------------------------------------------------------------
# Project-creation fixtures
# --------------------------------------------------------------------------


@pytest.fixture
def fresh_project(tmp_projects_dir: Path) -> Iterator[tuple[str, Path, sqlite3.Connection]]:
    """Create an empty project; yield (slug, db_path, conn). Closes conn on teardown."""
    name = "test-proj"
    _project_id, db_path = create_project(name=name, notes="e2e fixture")
    conn = connect(db_path)
    try:
        yield name, db_path, conn
    finally:
        conn.close()


def _make_synthetic_docx(target: Path, *, paragraphs: list[str] | None = None) -> Path:
    """Write a tiny .docx with one Heading 1 + a few paragraphs.

    Kept under ~200 bytes of body text so ingest + chunking is sub-millisecond.
    """
    doc = Document()
    doc.add_heading("Sample Specification Section", level=1)
    body = paragraphs or [
        "Contractor shall supply and install one (1) air handling unit.",
        "All ductwork shall be sealed to SMACNA Class A.",
        "Coordinate penetrations with the structural engineer.",
    ]
    for p in body:
        doc.add_paragraph(p)
    doc.save(target)
    return target


@pytest.fixture
def synthetic_docx(tmp_path: Path) -> Path:
    """Materialise a tiny synthetic DOCX under tmp_path/sample.docx."""
    return _make_synthetic_docx(tmp_path / "sample.docx")


@pytest.fixture
def fresh_project_with_sample_doc(
    fresh_project: tuple[str, Path, sqlite3.Connection],
    synthetic_docx: Path,
) -> tuple[str, Path, sqlite3.Connection, str]:
    """Create a project, ingest the synthetic doc, return (slug, db_path, conn, source_id)."""
    from meridian.ingest import ingest_file

    name, db_path, conn = fresh_project
    result = ingest_file(conn, file_path=synthetic_docx, project_root=synthetic_docx.parent)
    return name, db_path, conn, result.source_id


@pytest.fixture
def project_with_two_sources(
    fresh_project: tuple[str, Path, sqlite3.Connection],
    tmp_path: Path,
) -> tuple[sqlite3.Connection, list[str], list[str]]:
    """Create a project, ingest two synthetic documents, return (conn, source_ids, filenames).

    Used by tests that iterate over multiple sources (e.g., on_source_complete callback tests).
    """
    from meridian.ingest import ingest_file

    name, db_path, conn = fresh_project

    # Create and ingest two distinct synthetic documents
    doc1_path = _make_synthetic_docx(tmp_path / "doc1.docx", paragraphs=[
        "First document paragraph one.",
        "First document paragraph two.",
    ])
    doc2_path = _make_synthetic_docx(tmp_path / "doc2.docx", paragraphs=[
        "Second document paragraph one.",
        "Second document paragraph two.",
    ])

    result1 = ingest_file(conn, file_path=doc1_path, project_root=tmp_path)
    result2 = ingest_file(conn, file_path=doc2_path, project_root=tmp_path)

    source_ids = [result1.source_id, result2.source_id]
    filenames = ["doc1.docx", "doc2.docx"]

    return conn, source_ids, filenames


# --------------------------------------------------------------------------
# FastAPI test client
# --------------------------------------------------------------------------


@pytest.fixture
def fastapi_client(tmp_projects_dir: Path) -> Iterator[TestClient]:
    """Yield a Starlette TestClient bound to the live FastAPI app.

    Imports the app lazily so the ``tmp_projects_dir`` env-var override is
    in place before any module-level project-discovery code runs.
    """
    from meridian.api.main import app

    with TestClient(app) as client:
        yield client


@pytest.fixture
def api_client(fastapi_client: TestClient) -> TestClient:
    """Canonical alias used by task-4/5 specs; prefer this name in new tests."""
    return fastapi_client


# --------------------------------------------------------------------------
# LLM stubbing
# --------------------------------------------------------------------------


def _now() -> str:
    return datetime.now(UTC).strftime("%Y-%m-%dT%H:%M:%SZ")


def _persist_llm_row(
    conn: sqlite3.Connection,
    *,
    call_id: str,
    purpose: str,
    provider: str,
    model: str,
    system_prompt: str,
    user_prompt: str,
    response_text: str,
    parsed: Any | None,
    job_id: str | None,
) -> None:
    """Insert a stub llm_call row matching the schema the real client writes."""
    started_at = _now()
    finished_at = _now()
    with transaction(conn):
        conn.execute(
            """
            INSERT INTO llm_call (
                id, job_id, purpose, provider, model, provider_api_version,
                system_prompt, user_prompt, prompt_version_ref,
                temperature, top_p, max_tokens, input_hash,
                response_text, parsed_response,
                prompt_tokens, completion_tokens,
                cache_read_tokens, cache_write_tokens, cost_cents,
                started_at, finished_at, error
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                call_id,
                job_id,
                purpose,
                provider,
                model,
                None,
                system_prompt,
                user_prompt,
                f"{purpose}_v_test",
                0.0,
                None,
                16000,
                "deadbeef",
                response_text,
                json.dumps(parsed) if parsed is not None else None,
                10,
                10,
                None,
                None,
                0,
                started_at,
                finished_at,
                None,
            ),
        )


# Default canned responses keyed by ``purpose``. Tests can override entries
# (or pass a ``per_call`` callable) by setting attributes on the fixture
# object before triggering work.
DEFAULT_QUALITY_SCAN = {
    "scan_quality": "clean",
    "markups_present": False,
    "illegible_regions": [],
    "mismatched_references": [],
    "is_template": False,
    "document_class": "customer_requirements",
    "document_state": "100%",
    "revision": "A",
    "revision_detected_via": "filename_pattern",
    "is_demarcation_schedule": False,
    "extraction_path": "text_spec",
    "summary": "Synthetic test document.",
}

DEFAULT_TRIAGE_KEEP = {"keep": True, "reason": "stub: keep"}

# NB: trade/service/category values MUST match the locked default taxonomies
# (meridian.taxonomy.defaults) — using a value not in the seeded list would
# trip the ``taxonomy_new_value_proposed`` flag in persist.py and route the
# row to ``quarantined`` instead of ``auto_approved``, knocking it out of
# v_master_register and breaking downstream tender-list assertions.
DEFAULT_TEXT_SPEC = {
    "deliverables": [
        {
            "trade": "Mechanical",
            "service": "HVAC",
            "category": "design",
            "deliverables_summary": "Supply and install one (1) air handling unit.",
            "confidence": "high",
            "flags": [],
            "source_ref": "Section 1, paragraph 1",
            "applicable_standards": ["SMACNA Class A"],
        }
    ],
    "audit": [],
    "questions": [],
}


class StubLlm:
    """Mutable container for canned responses; exposed via the fixture object.

    Attributes (all overridable per-test):
        responses: dict[purpose -> response_dict]
        per_call: optional callable (purpose, system_prompt, user_prompt) -> dict
                  When set, takes precedence over ``responses``.
        recorded: list of (purpose, system_prompt, user_prompt) — for assertions
        fail_on: optional purpose to raise RuntimeError on (one-shot)
    """

    def __init__(self) -> None:
        self.responses: dict[str, dict[str, Any]] = {
            "quality_scan": dict(DEFAULT_QUALITY_SCAN),
            "triage": dict(DEFAULT_TRIAGE_KEEP),
            "extract_text_spec": dict(DEFAULT_TEXT_SPEC),
            "extract_demarcation": {"deliverables": [], "audit": [], "questions": []},
            "extract_bod": {"deliverables": [], "audit": [], "questions": []},
            "conflict_pass": {"conflicts": []},
            "error_explain": {"explanation": "stub"},
        }
        self.per_call: Callable[[str, str, str], dict[str, Any]] | None = None
        self.recorded: list[tuple[str, str, str]] = []
        self.fail_on: str | None = None
        self.extra_response_text: str = ""  # appended to response_text — used by
        # the secret-redaction test to inject a fake API key.

    def resolve(self, purpose: str, system_prompt: str, user_prompt: str) -> dict[str, Any]:
        if self.per_call is not None:
            return self.per_call(purpose, system_prompt, user_prompt)
        return self.responses.get(purpose, {})


@pytest.fixture
def mock_llm_client(monkeypatch: pytest.MonkeyPatch) -> StubLlm:
    """Replace ``meridian.llm.client.call_llm`` with a deterministic stub.

    The stub mirrors the real client's contract:
    - returns a populated ``LlmCall`` dataclass
    - inserts a row into the ``llm_call`` table (so FK references from
      ``deliverable.llm_call_id`` etc. resolve)
    - honours ``StubLlm.fail_on`` to simulate a mid-pipeline failure
    """
    stub = StubLlm()

    def _fake_call_llm(
        conn: sqlite3.Connection,
        *,
        purpose: str,
        provider: str | None = None,
        model: str | None = None,
        system_prompt: str,
        user_prompt: str,
        prompt_version_ref: str | None = None,
        job_id: str | None = None,
        temperature: float = 0.0,
        top_p: float | None = None,
        max_tokens: int = 16000,
        parse_json: bool = True,
    ) -> LlmCall:
        stub.recorded.append((purpose, system_prompt, user_prompt))
        if stub.fail_on == purpose:
            stub.fail_on = None  # one-shot
            raise RuntimeError(f"stub: forced failure for purpose={purpose}")

        parsed = stub.resolve(purpose, system_prompt, user_prompt)
        response_text = json.dumps(parsed) + stub.extra_response_text
        call_id = str(uuid.uuid4())
        _persist_llm_row(
            conn,
            call_id=call_id,
            purpose=purpose,
            provider=provider or "stub",
            model=model or "stub-model",
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            response_text=response_text,
            parsed=parsed,
            job_id=job_id,
        )
        return LlmCall(
            id=call_id,
            purpose=purpose,  # type: ignore[arg-type]
            response_text=response_text,
            parsed_json=parsed if parse_json else None,
            cost_cents=0,
            prompt_tokens=10,
            completion_tokens=10,
            cache_read_tokens=None,
            cache_write_tokens=None,
        )

    # Patch every call site that imported ``call_llm`` by name. Each
    # extract module did ``from meridian.llm.client import call_llm``, so
    # the binding lives on each module's namespace and must be patched
    # there as well as on the source module.
    monkeypatch.setattr("meridian.llm.client.call_llm", _fake_call_llm)
    for mod_path in (
        "meridian.extract.quality_scan",
        "meridian.extract.triage",
        "meridian.extract.text_spec",
        "meridian.extract.bod_import",
        "meridian.extract.demarcation",
        "meridian.extract.conflict_pass",
        "meridian.bootstrap.sweep",
    ):
        # Module hasn't been imported yet or doesn't bind the name — fine;
        # the source-level patch above still applies once the module
        # imports it the first time.
        with contextlib.suppress(AttributeError, ImportError):
            monkeypatch.setattr(f"{mod_path}.call_llm", _fake_call_llm)

    return stub


# --------------------------------------------------------------------------
# API-level project fixture (alpha-25)
# --------------------------------------------------------------------------


@pytest.fixture
def project_with_two_sources_via_api(fastapi_client, tmp_path) -> str:
    """Create a project + import two tiny DOCXs via the wizard's import-folder.

    Uses real DOCX files (not bare PDF bytes) because the ingest pipeline
    requires parseable file content. Returns the project slug so callers can
    drive the pipeline endpoint straight away.
    """
    folder = tmp_path / "alpha25_fixtures"
    folder.mkdir()
    _make_synthetic_docx(folder / "spec_a.docx", paragraphs=[
        "Contractor shall supply one (1) air handling unit.",
        "All ductwork sealed to SMACNA Class A.",
    ])
    _make_synthetic_docx(folder / "spec_b.docx", paragraphs=[
        "Electrical contractor shall install switchboard as shown.",
        "All cabling to AS/NZS 3000.",
    ])

    res = fastapi_client.post(
        "/api/setup/import-folder",
        json={"folder_path": str(folder), "project_name": "alpha25-fixture"},
    )
    assert res.status_code == 200, f"import-folder POST failed: {res.text}"
    job_id = res.json()["job_id"]

    deadline = time.monotonic() + 15.0
    last_status: dict = {}
    while time.monotonic() < deadline:
        s = fastapi_client.get(f"/api/setup/import-folder/{job_id}").json()
        last_status = s
        if s["status"] in {"succeeded", "failed"}:
            break
        time.sleep(0.05)

    assert last_status.get("status") == "succeeded", (
        f"import-folder job did not succeed within timeout: {last_status}"
    )
    return "alpha25-fixture"


# --------------------------------------------------------------------------
# Alpha-26: project_slug + async HTTP client fixtures
# --------------------------------------------------------------------------


@pytest.fixture
def project_slug(tmp_projects_dir: Path, fastapi_client: TestClient) -> str:
    """Create a project via the API and return its slug.

    Used by alpha-26 SSE tests that need a real project to exist before
    opening the events stream (the endpoint calls _ensure_project).
    """
    res = fastapi_client.post(
        "/api/projects",
        json={"name": "alpha26-fixture", "notes": "SSE e2e fixture"},
    )
    assert res.status_code == 200, res.text
    return "alpha26-fixture"


@pytest.fixture
async def api_client_async(fastapi_client: TestClient):
    """httpx.AsyncClient against the same FastAPI app for SSE streaming tests."""
    import httpx
    transport = httpx.ASGITransport(app=fastapi_client.app)
    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        yield client


# --------------------------------------------------------------------------
# Conflict-fixture (Task 7 — conflict_summary Excel column)
# --------------------------------------------------------------------------


@pytest.fixture
def project_with_conflicts(tmp_path):
    """In-memory project with one deliverable having a pending conflict.

    FK enforcement is disabled so we can insert rows without building the full
    extraction_job / llm_call prerequisite graph.  The fixture yields
    ``(conn, expected_pending_text)`` where ``expected_pending_text`` is the
    verbatim ``most_onerous_reasoning`` that should appear in the Excel column.
    """
    import json
    import sqlite3
    import uuid

    from meridian.db.connection import initialise

    db = tmp_path / "p.sqlite"
    conn = initialise(db)

    # Disable FK enforcement so we can insert conflict / deliverable rows
    # without building the full extraction_job + llm_call prerequisite graph.
    conn.execute("PRAGMA foreign_keys = OFF;")

    sd_id = str(uuid.uuid4())
    d_id = str(uuid.uuid4())
    c_id = str(uuid.uuid4())
    # Stub IDs for the FK columns we skip building.
    job_id = str(uuid.uuid4())
    llm_id = str(uuid.uuid4())
    group_id = str(uuid.uuid4())

    pending_text = (
        "The chiller-spec ≤22°C ceiling is not enforced in the global "
        "OSE document; the per-trade BoD is the more onerous source."
    )
    flag_label = f"conflicts_with_source_{c_id}"
    flag_context_json = json.dumps({flag_label: {"conflict_id": c_id}})
    flags_json = json.dumps([flag_label])

    conn.execute(
        """
        INSERT INTO source_document (
            id, filename, relative_path, content_hash, mime_type,
            size_bytes, imported_at, extraction_path
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (sd_id, "chiller-bod.pdf", "chiller-bod.pdf", f"hash_{sd_id}", "application/pdf",
         1, "2026-05-10T00:00:00Z", "text_spec"),
    )
    conn.execute(
        """
        INSERT INTO deliverable (
            id, extraction_group_id, source_id, extraction_job_id, llm_call_id,
            source_ref, applicable_standards, confidence, flags, flag_context,
            deliverables_summary, gate_outcome, status, auto_route_decision, created_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            d_id, group_id, sd_id, job_id, llm_id,
            '{"rendered": "Section 1"}', "[]", "high",
            flags_json, flag_context_json,
            "Chiller specification", "inside", "auto_approved", "auto_approved",
            "2026-05-10T00:00:00Z",
        ),
    )
    conn.execute(
        """
        INSERT INTO conflict (
            id, extraction_job_id, llm_call_id, kind,
            most_onerous_reasoning, status, created_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (c_id, job_id, llm_id, "cross_source_content",
         pending_text, "pending", "2026-05-10T00:00:00Z"),
    )
    conn.commit()

    yield conn, pending_text

    conn.execute("PRAGMA foreign_keys = ON;")
    conn.close()


# --------------------------------------------------------------------------
# Alpha-26 Task 6: conflict-register fixture (JSON + xlsx endpoints)
# --------------------------------------------------------------------------


@pytest.fixture
def project_with_conflicts_mixed_status(fastapi_client, tmp_projects_dir):
    """Create an API-reachable project with three conflicts (pending,
    resolved_accept_a, resolved_reject_both) so the conflict-register
    surface has multiple statuses to render.

    Unlike ``project_with_conflicts`` (which yields a raw SQLite connection
    unreachable by the API), this fixture creates the project via the API,
    then inserts conflict rows directly into the project DB via SQL so
    ``api_client`` can hit the /conflict-register endpoints.

    Yields the project slug string.
    """
    import uuid as _uuid

    from meridian.config import settings as _settings
    from meridian.db.connection import connect as _connect, initialise as _initialise
    from meridian.projects import project_db_path as _project_db_path

    slug = "alpha26-conflict-register"

    # Create the project via the API (so settings.projects_dir is populated
    # and _ensure_project can find the DB when the test hits the endpoints).
    res = fastapi_client.post(
        "/api/projects",
        json={"name": slug, "notes": "alpha-26 conflict register fixture"},
    )
    assert res.status_code == 200, f"project creation failed: {res.text}"

    db_path = _project_db_path(slug)
    assert db_path.exists(), f"DB not found at {db_path}"

    conn = _connect(db_path)
    # Disable FK + check enforcement so we can insert conflict rows without
    # building the full extraction_job / llm_call prerequisite graph.
    conn.execute("PRAGMA foreign_keys = OFF;")
    conn.execute("PRAGMA ignore_check_constraints = ON;")

    sd_id = str(_uuid.uuid4())
    job_id = str(_uuid.uuid4())
    llm_id = str(_uuid.uuid4())
    c_pending = str(_uuid.uuid4())
    c_resolved = str(_uuid.uuid4())
    c_reject = str(_uuid.uuid4())

    # Insert a stub source document so source-lookup joins return a filename.
    conn.execute(
        """
        INSERT INTO source_document (
            id, filename, relative_path, content_hash, mime_type,
            size_bytes, imported_at, extraction_path
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            sd_id, "fixture-source.pdf", "fixture-source.pdf",
            f"hash_{sd_id}", "application/pdf", 1,
            "2026-05-10T00:00:00Z", "text_spec",
        ),
    )

    # Three conflicts — pending, resolved_accept_a, resolved_reject_both.
    conn.execute(
        """
        INSERT INTO conflict (
            id, extraction_job_id, llm_call_id, kind,
            most_onerous_reasoning, status, created_at, resolved_at
        ) VALUES
            (?, ?, ?, 'cross_source_content',
             'The chiller spec ceiling is not enforced in the global OSE document.',
             'pending', '2026-05-10T00:00:00Z', NULL),
            (?, ?, ?, 'cross_source_content',
             'Accepted Source A as canonical; Source B value lower.',
             'resolved_accept_a', '2026-05-10T00:00:00Z', '2026-05-10T01:00:00Z'),
            (?, ?, ?, 'cross_source_content',
             'Both sources rejected; hybrid specification drafted.',
             'resolved_reject_both', '2026-05-10T00:00:00Z', '2026-05-10T02:00:00Z')
        """,
        (
            c_pending, job_id, llm_id,
            c_resolved, job_id, llm_id,
            c_reject, job_id, llm_id,
        ),
    )
    conn.commit()
    conn.execute("PRAGMA foreign_keys = ON;")
    conn.execute("PRAGMA ignore_check_constraints = OFF;")
    conn.close()

    yield slug
