"""Alpha-22 end-to-end: simulate the full wizard walk including a
projects_dir override at the project-creation step. Verify:
 - The named project DB exists at the chosen target_dir + slug.
 - That DB contains the rows imported during folder-import (NOT empty).
 - /api/setup/complete returns 200 (gates pass), NOT 400.

This is the regression test for the original bod-2 zero-sources bug.
"""

from __future__ import annotations

import sqlite3
import sys
import time
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient


def _make_synthetic_docx(target: Path) -> Path:
    """Write a tiny .docx so ingest_file can extract text without LLM."""
    doc = Document()
    doc.add_heading("Sample Specification Section", level=1)
    doc.add_paragraph("Contractor shall supply and install one (1) air handling unit.")
    doc.save(target)
    return target


def test_full_wizard_with_projects_dir_override_completes_successfully(
    fastapi_client: TestClient,
    tmp_path: Path,
    monkeypatch,
) -> None:
    # Hermetic wizard state — MERIDIAN_WIZARD_STATE_DIR is already set by
    # the ``tmp_projects_dir`` fixture (via ``fastapi_client``), but set it
    # explicitly here to a sub-dir so this test doesn't share state with
    # other tests that share the same tmp_path root.
    monkeypatch.setenv("MERIDIAN_WIZARD_STATE_DIR", str(tmp_path / "wizard_state"))

    # Reset rate-limit bucket and job registry from prior tests.
    from meridian.wizard import api as wizard_api
    wizard_api._rate_buckets.clear()
    with wizard_api._jobs_lock:
        wizard_api._jobs.clear()

    # Stub api-key validation so it returns "valid" without hitting Anthropic.
    # api.py imports validate_anthropic_key_str by name; patch that binding.
    monkeypatch.setattr(
        wizard_api,
        "validate_anthropic_key_str",
        lambda key: ("valid", "stubbed"),
    )

    # Stub keyring so the api-key step does not touch the OS keychain.
    class _FakeKeyring:
        _store: dict[tuple[str, str], str] = {}

        @staticmethod
        def set_password(service: str, user: str, value: str) -> None:
            _FakeKeyring._store[(service, user)] = value

        @staticmethod
        def get_password(service: str, user: str) -> str | None:
            return _FakeKeyring._store.get((service, user))

    monkeypatch.setitem(sys.modules, "keyring", _FakeKeyring)

    # 1. API key
    r = fastapi_client.post(
        "/api/setup/api-key",
        json={"key": "sk-ant-fake-key-not-real-just-for-test"},
    )
    assert r.status_code == 200, r.text

    # 2. Folder import — small folder with one synthetic .docx.
    import_src = tmp_path / "input"
    import_src.mkdir()
    _make_synthetic_docx(import_src / "sample.docx")

    r = fastapi_client.post(
        "/api/setup/import-folder",
        json={"folder_path": str(import_src), "project_name": "MYPROJ"},
    )
    assert r.status_code == 200, r.text
    job_id = r.json()["job_id"]

    # Poll until job completes.
    deadline = time.monotonic() + 30
    while time.monotonic() < deadline:
        r = fastapi_client.get(f"/api/setup/import-folder/{job_id}")
        assert r.status_code == 200, r.text
        if r.json()["status"] in ("succeeded", "failed"):
            break
        time.sleep(0.1)
    assert r.json()["status"] == "succeeded", r.json()

    # 3. Pick a DIFFERENT projects_dir at project-creation time.
    final_dir = tmp_path / "user-chosen-projects-dir"
    final_dir.mkdir()

    # Suggest-name.
    r = fastapi_client.post(
        "/api/setup/projects/suggest-name",
        json={"folder_path": str(import_src)},
    )
    assert r.status_code == 200, r.text
    suggested = r.json()["suggested_name"]

    r = fastapi_client.post(
        "/api/setup/projects",
        json={
            "name": suggested,
            "slug": suggested,
            "projects_dir": str(final_dir),
            "notes": None,
        },
    )
    assert r.status_code == 200, r.text

    # 4. Final DB must exist at user-chosen dir AND contain the imported source.
    from meridian.projects import _slugify
    slug = _slugify(suggested)
    final_db = final_dir / f"{slug}.sqlite"
    assert final_db.exists(), f"final DB missing at {final_db}"

    with sqlite3.connect(final_db) as conn:
        sources = list(conn.execute("SELECT * FROM source_document"))
        assert len(sources) == 1, (
            f"final DB has unexpected source count — staging DB was not adopted correctly. "
            f"sources={sources}"
        )

    # 5. /api/setup/complete must succeed (gates satisfied).
    r = fastapi_client.post("/api/setup/complete")
    assert r.status_code == 200, (
        f"setup/complete returned {r.status_code}: {r.text} — "
        "wizard state is inconsistent across projects_dir change."
    )


def test_setup_create_project_case_2_same_path_same_slug(
    fastapi_client: TestClient,
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Case 2: user accepts the auto-derived slug AND the default
    projects_dir. Should run an in-place UPDATE of project.name without
    moving the file or minting a fresh DB."""
    monkeypatch.setenv("MERIDIAN_WIZARD_STATE_DIR", str(tmp_path / "wizard_state"))
    from meridian.config import settings
    monkeypatch.setattr(settings, "data_dir", tmp_path / "projects_dir")
    (tmp_path / "projects_dir").mkdir()

    from meridian.wizard import api as wizard_api
    monkeypatch.setattr(
        wizard_api, "validate_anthropic_key_str", lambda key: ("valid", "stubbed")
    )

    r = fastapi_client.post(
        "/api/setup/api-key",
        json={"key": "sk-ant-fake-key-not-real-just-for-test"},
    )
    assert r.status_code == 200, r.text

    import_src = tmp_path / "BOD"
    import_src.mkdir()
    _make_synthetic_docx(import_src / "sample.docx")

    r = fastapi_client.post(
        "/api/setup/import-folder",
        json={"folder_path": str(import_src), "project_name": "BOD"},
    )
    assert r.status_code == 200, r.text
    job_id = r.json()["job_id"]

    deadline = time.monotonic() + 30
    while time.monotonic() < deadline:
        r = fastapi_client.get(f"/api/setup/import-folder/{job_id}")
        if r.json()["status"] in ("succeeded", "failed"):
            break
        time.sleep(0.1)
    assert r.json()["status"] == "succeeded", r.json()

    # Now post /api/setup/projects with SAME slug AND SAME projects_dir.
    # This exercises case 2.
    r = fastapi_client.post(
        "/api/setup/projects",
        json={
            "name": "Building of Dreams",  # different display name
            "slug": "bod",  # same slug as auto-derived from "BOD" folder
            "projects_dir": str(tmp_path / "projects_dir"),
            "notes": None,
        },
    )
    assert r.status_code == 200, r.text

    # The staging DB should still be at the SAME path (no move happened).
    db_at_default = (tmp_path / "projects_dir") / "bod.sqlite"
    assert db_at_default.exists()

    # And its project.name should be "Building of Dreams" (the in-place UPDATE).
    with sqlite3.connect(db_at_default) as conn:
        rows = list(conn.execute("SELECT name FROM project"))
        assert rows == [("Building of Dreams",)], rows


def test_setup_create_project_case_3_same_path_different_slug(
    fastapi_client: TestClient,
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Case 3: user keeps the default projects_dir but RENAMES the project
    to a different slug. Should adopt_project (rename + name update) within
    the same dir."""
    monkeypatch.setenv("MERIDIAN_WIZARD_STATE_DIR", str(tmp_path / "wizard_state"))
    from meridian.config import settings
    monkeypatch.setattr(settings, "data_dir", tmp_path / "projects_dir")
    (tmp_path / "projects_dir").mkdir()

    from meridian.wizard import api as wizard_api
    monkeypatch.setattr(
        wizard_api, "validate_anthropic_key_str", lambda key: ("valid", "stubbed")
    )

    r = fastapi_client.post(
        "/api/setup/api-key",
        json={"key": "sk-ant-fake-key-not-real-just-for-test"},
    )
    assert r.status_code == 200, r.text

    import_src = tmp_path / "BOD"
    import_src.mkdir()
    _make_synthetic_docx(import_src / "sample.docx")

    r = fastapi_client.post(
        "/api/setup/import-folder",
        json={"folder_path": str(import_src), "project_name": "BOD"},
    )
    assert r.status_code == 200, r.text
    job_id = r.json()["job_id"]

    deadline = time.monotonic() + 30
    while time.monotonic() < deadline:
        r = fastapi_client.get(f"/api/setup/import-folder/{job_id}")
        if r.json()["status"] in ("succeeded", "failed"):
            break
        time.sleep(0.1)
    assert r.json()["status"] == "succeeded", r.json()

    # Now post /api/setup/projects with SAME projects_dir but DIFFERENT slug.
    r = fastapi_client.post(
        "/api/setup/projects",
        json={
            "name": "MyProject",
            "slug": "my-project",
            "projects_dir": str(tmp_path / "projects_dir"),
            "notes": None,
        },
    )
    assert r.status_code == 200, r.text

    # The staging DB at "bod.sqlite" should be GONE — moved to "my-project.sqlite".
    old_db = (tmp_path / "projects_dir") / "bod.sqlite"
    new_db = (tmp_path / "projects_dir") / "my-project.sqlite"
    assert not old_db.exists(), "staging bod.sqlite should have been moved"
    assert new_db.exists(), "new my-project.sqlite should exist"

    # And the source rows should be in the new DB.
    with sqlite3.connect(new_db) as conn:
        sources = list(conn.execute("SELECT * FROM source_document"))
        assert len(sources) == 1, sources
